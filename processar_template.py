#!/usr/bin/env python3
"""
Script para processar templates DOCX com python-docx
Preenche campos do template com dados do JSON
"""

import sys
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html.parser import HTMLParser
from io import StringIO

class HTMLToDocxConverter(HTMLParser):
    """Converte HTML sanitizado para elementos do python-docx"""
    
    def __init__(self, document):
        super().__init__()
        self.document = document
        self.current_paragraph = None
        self.current_run = None
        self.in_bold = False
        self.in_italic = False
        self.in_underline = False
        self.in_heading = 0
        self.list_level = 0
        
    def handle_starttag(self, tag, attrs):
        if tag == 'p':
            self.current_paragraph = self.document.add_paragraph()
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            self.in_heading = level
            self.current_paragraph = self.document.add_heading('', level=level)
        elif tag == 'b' or tag == 'strong':
            self.in_bold = True
        elif tag == 'i' or tag == 'em':
            self.in_italic = True
        elif tag == 'u':
            self.in_underline = True
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')
        elif tag == 'ul' or tag == 'ol':
            self.list_level += 1
        elif tag == 'li':
            self.current_paragraph = self.document.add_paragraph(style='List Bullet' if self.list_level > 0 else 'Normal')
        elif tag == 'blockquote':
            self.current_paragraph = self.document.add_paragraph()
            self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
            self.current_paragraph.paragraph_format.right_indent = Inches(0.5)
    
    def handle_endtag(self, tag):
        if tag == 'p' or tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.current_paragraph = None
            self.in_heading = 0
        elif tag == 'b' or tag == 'strong':
            self.in_bold = False
        elif tag == 'i' or tag == 'em':
            self.in_italic = False
        elif tag == 'u':
            self.in_underline = False
        elif tag == 'ul' or tag == 'ol':
            self.list_level -= 1
    
    def handle_data(self, data):
        if not self.current_paragraph:
            self.current_paragraph = self.document.add_paragraph()
        
        run = self.current_paragraph.add_run(data)
        
        if self.in_bold:
            run.bold = True
        if self.in_italic:
            run.italic = True
        if self.in_underline:
            run.underline = True

def substituir_campos(doc, dados):
    """
    Substitui campos {{variavel}} no documento pelos valores do JSON
    """
    # Substituir em parágrafos
    for paragraph in doc.paragraphs:
        for key, value in dados.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                # Substituir mantendo formatação
                for run in paragraph.runs:
                    if f'{{{{{key}}}}}' in run.text:
                        run.text = run.text.replace(f'{{{{{key}}}}}', str(value))
    
    # Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dados.items():
                    if f'{{{{{key}}}}}' in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if f'{{{{{key}}}}}' in run.text:
                                    run.text = run.text.replace(f'{{{{{key}}}}}', str(value))
    
    # Substituir em cabeçalhos e rodapés
    for section in doc.sections:
        # Cabeçalho
        for paragraph in section.header.paragraphs:
            for key, value in dados.items():
                if f'{{{{{key}}}}}' in paragraph.text:
                    for run in paragraph.runs:
                        if f'{{{{{key}}}}}' in run.text:
                            run.text = run.text.replace(f'{{{{{key}}}}}', str(value))
        
        # Rodapé
        for paragraph in section.footer.paragraphs:
            for key, value in dados.items():
                if f'{{{{{key}}}}}' in paragraph.text:
                    for run in paragraph.runs:
                        if f'{{{{{key}}}}}' in run.text:
                            run.text = run.text.replace(f'{{{{{key}}}}}', str(value))

def inserir_conteudo_html(doc, html_content, posicao_marcador='{{conteudo}}'):
    """
    Insere conteúdo HTML convertido no lugar do marcador
    """
    # Encontrar parágrafo com marcador
    marcador_encontrado = False
    indice_paragrafo = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        if posicao_marcador in paragraph.text:
            marcador_encontrado = True
            indice_paragrafo = i
            # Remover marcador
            paragraph.text = paragraph.text.replace(posicao_marcador, '')
            break
    
    if not marcador_encontrado:
        print(f"⚠️  Marcador {posicao_marcador} não encontrado, adicionando no final")
        indice_paragrafo = len(doc.paragraphs)
    
    # Converter HTML para elementos do documento
    # Criar documento temporário para conversão
    temp_doc = Document()
    converter = HTMLToDocxConverter(temp_doc)
    converter.feed(html_content)
    
    # Inserir parágrafos convertidos no documento principal
    # (Nota: python-docx não tem método direto para inserir em posição específica,
    # então adicionamos no final. Para posicionamento preciso, seria necessário
    # manipular o XML subjacente)
    for paragraph in temp_doc.paragraphs:
        new_p = doc.add_paragraph(paragraph.text, style=paragraph.style)
        new_p.alignment = paragraph.alignment
        
        # Copiar formatação dos runs
        for i, run in enumerate(paragraph.runs):
            if i < len(new_p.runs):
                new_run = new_p.runs[i]
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

def processar_template(template_path, dados_path, output_path):
    """
    Processa template DOCX preenchendo com dados do JSON
    
    Args:
        template_path: Caminho para o template .docx
        dados_path: Caminho para o arquivo JSON com dados
        output_path: Caminho para salvar o documento final
    """
    
    print(f"📄 Template: {template_path}")
    print(f"📊 Dados: {dados_path}")
    print(f"💾 Output: {output_path}")
    
    # Carregar dados
    with open(dados_path, 'r', encoding='utf-8') as f:
        dados = json.load(f)
    
    print(f"✅ Dados carregados: {len(dados)} campos")
    
    # Abrir template
    try:
        doc = Document(template_path)
        print(f"✅ Template aberto: {len(doc.paragraphs)} parágrafos")
    except Exception as e:
        print(f"❌ Erro ao abrir template: {e}")
        raise
    
    # Preparar dados para substituição
    dados_substituicao = {
        'titulo': dados.get('titulo', ''),
        'autor': dados.get('autor', ''),
        'orientador': dados.get('orientador', ''),
        'coorientador': dados.get('coorientador', ''),
        'instituicao': dados.get('instituicao', ''),
        'programa': dados.get('programa', ''),
        'departamento': dados.get('departamento', ''),
        'cidade': dados.get('cidade', ''),
        'ano': dados.get('ano', ''),
        'data_completa': dados.get('data_completa', ''),
        'tipo_trabalho': dados.get('tipo_trabalho', ''),
        'resumo': dados.get('resumo', ''),
        'abstract': dados.get('abstract', ''),
        'palavras_chave': dados.get('palavras_chave', ''),
        'keywords': dados.get('keywords', ''),
    }
    
    # Substituir campos básicos
    print("🔄 Substituindo campos básicos...")
    substituir_campos(doc, dados_substituicao)
    
    # Inserir conteúdo HTML se existir
    if 'estrutura' in dados and isinstance(dados['estrutura'], list):
        print(f"🔄 Processando {len(dados['estrutura'])} seções...")
        
        # Concatenar conteúdo de todas as seções
        conteudo_completo = []
        for secao in dados['estrutura']:
            if 'conteudo' in secao:
                conteudo_completo.append(secao['conteudo'])
        
        html_conteudo = '\n'.join(conteudo_completo)
        
        if html_conteudo:
            print(f"🔄 Inserindo conteúdo HTML ({len(html_conteudo)} caracteres)...")
            inserir_conteudo_html(doc, html_conteudo)
    
    # Inserir referências se existirem
    if 'referencias' in dados:
        print("🔄 Inserindo referências...")
        inserir_conteudo_html(doc, dados['referencias'], '{{referencias}}')
    
    # Salvar documento processado
    try:
        doc.save(output_path)
        print(f"✅ Documento salvo: {output_path}")
    except Exception as e:
        print(f"❌ Erro ao salvar documento: {e}")
        raise
    
    # Verificar se arquivo foi criado
    if not os.path.exists(output_path):
        raise Exception("Arquivo de saída não foi criado")
    
    # Verificar tamanho do arquivo
    size = os.path.getsize(output_path)
    print(f"📦 Tamanho do arquivo: {size} bytes")
    
    if size < 1000:
        print("⚠️  Aviso: Arquivo muito pequeno, pode estar vazio")
    
    return True

def main():
    if len(sys.argv) != 4:
        print("Uso: python3 processar_template.py <template.docx> <dados.json> <output.docx>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    dados_path = sys.argv[2]
    output_path = sys.argv[3]
    
    # Validar arquivos de entrada
    if not os.path.exists(template_path):
        print(f"❌ Erro: Template não encontrado: {template_path}")
        sys.exit(1)
    
    if not os.path.exists(dados_path):
        print(f"❌ Erro: Arquivo de dados não encontrado: {dados_path}")
        sys.exit(1)
    
    try:
        sucesso = processar_template(template_path, dados_path, output_path)
        if sucesso:
            print("✅ Processamento concluído com sucesso!")
            sys.exit(0)
        else:
            print("❌ Erro no processamento")
            sys.exit(1)
    except Exception as e:
        print(f"❌ Erro fatal: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
